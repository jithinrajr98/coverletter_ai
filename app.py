import streamlit as st
from datetime import datetime
from llm_utils import cover_letter_prompt, profile_modifier_prompt, job_analysis_prompt
from docx import Document
import os
from docx2pdf import convert
import tempfile
from fpdf import FPDF
import google.generativeai as genai
from groq import Groq
from config.styles import apply_custom_styles, set_page_config, header_section
from dotenv import load_dotenv
load_dotenv()


# Configure APIs

# configure Groq
api_key = None
if hasattr(st, 'secrets') and 'GROQ_API_KEY' in st.secrets:
    api_key = st.secrets['GROQ_API_KEY']
elif 'GROQ_API_KEY' in os.environ:
    api_key = os.environ['GROQ_API_KEY']
if not api_key:
    raise ValueError("GROQ_API_KEY not found in Streamlit secrets or environment variables.")

GROQ_CLIENT = Groq(api_key=api_key)

# Models
groq_model = "meta-llama/llama-4-scout-17b-16e-instruct"

# Initialize session state
def init_session_state():
    defaults = {
        'cover_letter': "",
        'modified_resume': "",
        'execution_time': None,
        'cover_letter_fr': "",
        'modified_resume_fr': "",
        'show_french': False,
        'show_french_profile': False,
        'job_analysis': ""
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

# Default profile text
DEFAULT_PROFILE = """AI Engineer with 5+ years of experience deploying production-grade ML/Deep learning models and multi-agent systems. Led a preventive maintenance project for a Fortune 500 client, achieving $1.8M/year in savings. Expert in building agentic AI workflows, RAG architectures, and cloud-native ML pipelines."""



def translate_to_french(text):
    """Translate text to French using Groq"""
    try:
        response = GROQ_CLIENT.chat.completions.create(
            messages=[{
                "role": "user", 
                "content": f"Translate this professional text to French. Maintain the professional tone and formatting:\n\n{text}"
            }],
            model=groq_model
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Translation error: {str(e)}")
        return None

def update_cover_letter_state():
    """Update cover letter state when text area is edited"""
    if st.session_state.show_french:
        # Update French version
        if f"cover_letter_display_fr" in st.session_state:
            st.session_state.cover_letter_fr = st.session_state.cover_letter_display_fr
    else:
        # Update English version
        if f"cover_letter_display_en" in st.session_state:
            st.session_state.cover_letter = st.session_state.cover_letter_display_en

def update_profile_section(doc_path, new_profile_text):
    """Update the profile section in a Word document"""
    try:
        doc = Document(doc_path)
        profile_updated = False
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip().upper() in ["PROFILE", "SUMMARY", "PROFIL"]:
                if i + 1 < len(doc.paragraphs):
                    doc.paragraphs[i+1].clear()
                    doc.paragraphs[i+1].add_run(new_profile_text)
                    profile_updated = True
                    break
        
        if profile_updated:
            doc.save(doc_path)
            return True
        return False
    except Exception as e:
        st.error(f"Error updating DOCX file: {str(e)}")
        return False

def convert_to_pdf(docx_path):
    """Convert DOCX to PDF and return PDF path"""
    try:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_pdf:
            pdf_path = temp_pdf.name
        convert(docx_path, pdf_path)
        return pdf_path
    except Exception as e:
        st.error(f"Error converting to PDF: {str(e)}")
        return None

def create_pdf_from_text(content, filename):
    """Create a PDF from text content with better formatting"""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=11)
        pdf.set_margins(20, 20, 20)
        
        # Split content into lines and handle text wrapping
        lines = content.split('\n')
        for line in lines:
            if line.strip():  # Skip empty lines
                # Handle long lines by splitting them
                words = line.split(' ')
                current_line = ""
                for word in words:
                    test_line = current_line + " " + word if current_line else word
                    if pdf.get_string_width(test_line) < 170:  # Max width
                        current_line = test_line
                    else:
                        if current_line:
                            pdf.cell(0, 6, txt=current_line.encode('latin-1', 'replace').decode('latin-1'), ln=1)
                        current_line = word
                if current_line:
                    pdf.cell(0, 6, txt=current_line.encode('latin-1', 'replace').decode('latin-1'), ln=1)
            else:
                pdf.cell(0, 6, txt="", ln=1)  # Empty line
        
        temp_pdf_path = f"{filename}.pdf"
        pdf.output(temp_pdf_path)
        return temp_pdf_path
    except Exception as e:
        st.error(f"Error creating PDF: {str(e)}")
        return None

def load_resume():
    """Load resume content from file"""
    try:
        with open("my_resume.txt", 'r', encoding='utf-8') as file:
            resume_text = file.read()
        
        if not st.session_state.modified_resume:
            st.session_state.modified_resume = DEFAULT_PROFILE
            
        return resume_text
    except FileNotFoundError:
        st.error("Resume file not found. Please ensure 'my_resume.txt' exists.")
        return ""
    except Exception as e:
        st.error(f"Error loading resume: {str(e)}")
        return ""

def main():
    # Initialize
    init_session_state()
    
    # Set page config
    apply_custom_styles()
    set_page_config ()
    header_section()
    
    st.divider()

    # Apply dark theme
    
    # Load resume
    resume = load_resume()
    
 
    
    # Main layout
    col1, col2 = st.columns([1.2, 1])
    
    with col1:
        st.markdown("### 📋 Job Description")
        job_description = st.text_area(
            "Paste the complete job description:",
            height=300,
            placeholder="Paste the complete job description including requirements and company information...",
            help="Include all relevant details for better customization"
        )
        
        # Action buttons
        col_btn1, col_btn2, col_btn3 = st.columns([1, 1, 1])
        st.divider()

        with col_btn1:
            if st.button("🔍 Analyze Job", type="secondary", use_container_width=True):
                if not job_description.strip():
                    st.warning("⚠️ Please enter a job description first.")
                else:
                    with st.spinner("🔄 Analyzing job description..."):
                        try:
                            start_time = datetime.now()
                            response = GROQ_CLIENT.chat.completions.create(
                                messages=[{"role": "user", "content": job_analysis_prompt(job_description)}],
                                model=groq_model
                            )
                            st.session_state.job_analysis = response.choices[0].message.content
                            st.session_state.execution_time = datetime.now() - start_time
                            st.success("✅ Job analysis completed!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ Error: {str(e)}")
        
        with col_btn2:
            if st.button("✨ Generate Cover Letter", type="primary", use_container_width=True):
                if not job_description.strip():
                    st.warning("⚠️ Please enter a job description first.")
                elif not resume:
                    st.error("❌ Resume content is missing.")
                else:
                    with st.spinner("🔄 Generating your personalized cover letter..."):
                        try:
                            start_time = datetime.now()
                            response = GROQ_CLIENT.chat.completions.create(
                                messages=[{"role": "user", "content": cover_letter_prompt(job_description, resume)}],
                                model=groq_model
                            )
                            st.session_state.cover_letter = response.choices[0].message.content
                            st.session_state.execution_time = datetime.now() - start_time
                            st.session_state.show_french = False  # Reset French view
                            st.success("✅ Cover letter generated successfully!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ Error: {str(e)}")
        
        with col_btn3:
            if st.button("🎯 Optimize CV Profile", type="secondary", use_container_width=True):
                if not job_description.strip():
                    st.warning("⚠️ Please enter a job description first.")
                elif not resume:
                    st.error("❌ Resume content is missing.")
                else:
                    with st.spinner("🔄 Tailoring your CV profile..."):
                        try:
                            start_time = datetime.now()
                            response = GROQ_CLIENT.chat.completions.create(
                                messages=[{"role": "user", "content": profile_modifier_prompt(job_description, st.session_state.modified_resume)}],
                                model=groq_model
                            )
                            st.session_state.modified_resume = response.choices[0].message.content
                            st.session_state.modified_resume_fr = ""  # Reset French translation
                            st.session_state.show_french_profile = False  # Reset to English view
                            st.session_state.execution_time = datetime.now() - start_time
                            st.success("✅ CV profile optimized successfully!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ Error: {str(e)}")
        
        # Job Analysis Display
        if st.session_state.job_analysis:
            st.markdown("### 📊 Job Analysis")
                # st.text_area(
                #     "Job Analysis Results:",
                #     value=st.session_state.job_analysis,
                #     height=500,
                #     key="job_analysis_display"
                # )
                
            st.success(st.session_state.job_analysis)
            
            # Download job analysis
            st.download_button(
                label="📄 Download Analysis",
                data=st.session_state.job_analysis,
                file_name=f"job_analysis_{datetime.now().strftime('%Y%m%d')}.txt",
                mime="text/plain",
                use_container_width=True
            )
    
    with col2:
        tab1, tab2 = st.tabs(["📝 Cover Letter", "👤 CV Profile"])
        
        with tab1:
            if st.session_state.cover_letter:
                # Language toggle buttons
                col_en, col_fr = st.columns(2)
                with col_en:
                    if st.button("🇺🇸 English", use_container_width=True, type="primary" if not st.session_state.show_french else "secondary"):
                        st.session_state.show_french = False
                        st.rerun()
                
                with col_fr:
                    if st.button("🇫🇷 Français", use_container_width=True, type="primary" if st.session_state.show_french else "secondary"):
                        # Always translate based on current English content
                        with st.spinner("🔄 Translating to French..."):
                            # Get current value from English text area
                            current_english = st.session_state.get(f"cover_letter_display_en", st.session_state.cover_letter)
                            st.session_state.cover_letter_fr = translate_to_french(current_english)
                        st.session_state.show_french = True
                        st.rerun()
                
                # Display appropriate version
                current_letter = st.session_state.cover_letter_fr if st.session_state.show_french and st.session_state.cover_letter_fr else st.session_state.cover_letter
                language_suffix = "_fr" if st.session_state.show_french else "_en"
                
                st.text_area(
                    f"Your cover letter ({'French' if st.session_state.show_french else 'English'}):",
                    value=current_letter,
                    height=500,
                    key=f"cover_letter_display{language_suffix}",
                    on_change=lambda: update_cover_letter_state()
                )
                
                # Download buttons
                col_txt, col_pdf = st.columns(2)
                with col_txt:
                    st.download_button(
                        label="📄 Download TXT",
                        data=current_letter,
                        file_name=f"cover_letter_{datetime.now().strftime('%Y%m%d')}{language_suffix}.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                
                with col_pdf:
                    if st.button("📊 Generate PDF", use_container_width=True):
                        with st.spinner("📄 Creating PDF..."):
                            pdf_path = create_pdf_from_text(
                                current_letter,
                                f"Jithin_Reghuvaran_coverletter{language_suffix}"
                            )
                            if pdf_path:
                                with open(pdf_path, "rb") as f:
                                    st.download_button(
                                        label="⬇️ Download PDF",
                                        data=f,
                                        file_name=f"Jithin_Reghuvaran_coverletter{language_suffix}.pdf",
                                        mime="application/pdf",
                                        use_container_width=True
                                    )
                                os.remove(pdf_path)
            else:
                st.info("👆 Generate a cover letter using the job description")
        
        with tab2:
            # Language toggle buttons for profile
            col_en_prof, col_fr_prof = st.columns(2)
            with col_en_prof:
                if st.button("🇺🇸 English Profile", use_container_width=True, type="primary" if not st.session_state.show_french_profile else "secondary"):
                    st.session_state.show_french_profile = False
                    st.rerun()
            
            with col_fr_prof:
                if st.button("🇫🇷 Profil Français", use_container_width=True, type="primary" if st.session_state.show_french_profile else "secondary"):
                    if not st.session_state.modified_resume_fr:
                        with st.spinner("🔄 Translating profile to French..."):
                            st.session_state.modified_resume_fr = translate_to_french(st.session_state.modified_resume)
                    st.session_state.show_french_profile = True
                    st.rerun()
            
            # Display appropriate profile version
            current_profile = st.session_state.modified_resume_fr if st.session_state.show_french_profile and st.session_state.modified_resume_fr else st.session_state.modified_resume
            profile_suffix = "_fr" if st.session_state.show_french_profile else "_en"
            
            modified_profile = st.text_area(
                f"Edit your profile section ({'French' if st.session_state.show_french_profile else 'English'}):",
                value=current_profile,
                height=300,
                help="This will be used to update your CV profile section",
                key=f"profile_editor{profile_suffix}"
            )
            
            # Update the appropriate session state when user edits
            if st.session_state.show_french_profile:
                if modified_profile != st.session_state.modified_resume_fr:
                    st.session_state.modified_resume_fr = modified_profile
            else:
                if modified_profile != st.session_state.modified_resume:
                    st.session_state.modified_resume = modified_profile
            
            # Download profile buttons
            col_prof_txt, col_prof_pdf = st.columns(2)
            with col_prof_txt:
                st.download_button(
                    label="📄 Download Profile TXT",
                    data=current_profile,
                    file_name=f"profile_summary_{datetime.now().strftime('%Y%m%d')}{profile_suffix}.txt",
                    mime="text/plain",
                    use_container_width=True
                )
            
            with col_prof_pdf:
                if st.button("📊 Generate Profile PDF", use_container_width=True):
                    with st.spinner("📄 Creating Profile PDF..."):
                        pdf_path = create_pdf_from_text(
                            current_profile,
                            f"Jithin_Reghuvaran_profile{profile_suffix}"
                        )
                        if pdf_path:
                            with open(pdf_path, "rb") as f:
                                st.download_button(
                                    label="⬇️ Download Profile PDF",
                                    data=f,
                                    file_name=f"Jithin_Reghuvaran_profile{profile_suffix}.pdf",
                                    mime="application/pdf",
                                    use_container_width=True,
                                    key=f"profile_pdf_download{profile_suffix}"
                                )
                            os.remove(pdf_path)
            
            # DOCX Update Section
            st.markdown("### 📄 Update Word Document")
            docx_file = st.file_uploader("Upload your CV (DOCX)", type=["docx"], help="Upload your CV to update the profile section")
            
            if docx_file is not None:
                # Option to choose language for DOCX update
                update_lang = st.radio(
                    "Choose language for CV update:",
                    ["English", "French"],
                    horizontal=True,
                    help="Select which version of the profile to use for updating your CV"
                )
                
                if st.button("🔄 Update DOCX Profile", use_container_width=True):
                    with st.spinner("📄 Updating and converting your document..."):
                        try:
                            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
                                temp_docx_path = temp_docx.name
                                temp_docx.write(docx_file.getbuffer())
                            
                            # Choose which profile version to use
                            profile_to_use = st.session_state.modified_resume_fr if update_lang == "French" and st.session_state.modified_resume_fr else st.session_state.modified_resume
                            lang_suffix = "_fr" if update_lang == "French" else "_en"
                            
                            success = update_profile_section(
                                temp_docx_path,
                                profile_to_use.replace("Profile:\n", "").replace("Profil:\n", "")
                            )
                            
                            if success:
                                pdf_path = convert_to_pdf(temp_docx_path)
                                
                                if pdf_path:
                                    with open(pdf_path, "rb") as f:
                                        st.download_button(
                                            label=f"⬇️ Download Updated CV ({update_lang}) PDF",
                                            data=f,
                                            file_name=f"Jithin_Reghuvaran_CV{lang_suffix}.pdf",
                                            mime="application/pdf",
                                            use_container_width=True,
                                            key=f"cv_download{lang_suffix}"
                                        )
                                    st.success("✅ PDF generated successfully!")
                                    os.remove(pdf_path)
                                else:
                                    st.warning("⚠️ Could not convert to PDF")
                            else:
                                st.warning("⚠️ Profile section not found in document")
                        except Exception as e:
                            st.error(f"❌ Error processing document: {str(e)}")
                        finally:
                            if os.path.exists(temp_docx_path):
                                os.remove(temp_docx_path)
    
    # Execution time display
    if st.session_state.execution_time:
        st.info(f"⏱️ Execution Time: {st.session_state.execution_time}")

if __name__ == "__main__":
    main()